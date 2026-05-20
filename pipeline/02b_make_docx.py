"""Step 1.5: aggregate per-song lyric .txt files into a single .docx archive.

Reads:  data/lyrics/{raw|en}/*.txt
Writes: data/lyrics/aggregated/{batch}{_en}.docx

This is purely an archive step — downstream pipeline doesn't read the .docx.

Run:
    python pipeline/02b_make_docx.py --batch 2026-05-17                  # raw zh
    python pipeline/02b_make_docx.py --batch 2026-05-17 --lang en        # translated en
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
from pipeline._lib import paths  # noqa: E402

SEPARATOR = "=" * 60


def build(lang: str, batch: str, songs: list[str] | None) -> Path:
    src_dir = paths.LYRICS_RAW if lang == "zh" else paths.LYRICS_EN
    if not src_dir.exists():
        raise FileNotFoundError(src_dir)

    if songs:
        files = [src_dir / f"{s}.txt" for s in songs]
        missing = [f for f in files if not f.exists()]
        if missing:
            raise FileNotFoundError(f"missing: {missing}")
    else:
        files = sorted(src_dir.glob("*.txt"))
    if not files:
        raise RuntimeError(f"no lyric files in {src_dir}")

    doc = Document()
    doc.add_heading(f"Lyrics archive — {batch} ({lang})", level=1)
    for f in files:
        doc.add_paragraph(SEPARATOR)
        doc.add_heading(f.stem, level=2)
        text = f.read_text(encoding="utf-8")
        for line in text.splitlines():
            doc.add_paragraph(line)

    paths.LYRICS_AGG.mkdir(parents=True, exist_ok=True)
    suffix = "_en" if lang == "en" else ""
    out = paths.LYRICS_AGG / f"{batch}{suffix}.docx"
    doc.save(out)
    return out


def main() -> int:
    p = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    p.add_argument("--batch", required=True)
    p.add_argument("--lang", choices=["zh", "en"], default="zh")
    p.add_argument("--song", action="append", default=[], help="limit to these songs; default: all")
    args = p.parse_args()

    out = build(args.lang, args.batch, args.song or None)
    print(f"wrote {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
